import os
import json
import tempfile
from datetime import datetime

import requests
import streamlit as st
import openpyxl
from docx import Document


# =========================
# CONFIG
# =========================
OLLAMA_URL = "http://localhost:11434/api/generate"
DEFAULT_MODEL = "llama3.1:8b"
FALLBACK_MODEL = "phi3:mini"
DEFAULT_SHEET_NAME = "Checklist"

TRUE_LIKE = {"true", "yes", "y", "1", "x", "✓", "✔", "checked"}


# =========================
# HELPERS
# =========================
def is_ticked(value) -> bool:
    if value is None:
        return False
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        return value.strip().lower() in TRUE_LIKE
    return False


def safe_str(value) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    return str(value).strip()


def read_qa_block(ws, start_row: int, end_row: int, q_col: int = 2, a_col: int = 4):
    items = []
    for r in range(start_row, end_row + 1):
        criteria = safe_str(ws.cell(r, q_col).value)
        commentary = safe_str(ws.cell(r, a_col).value)

        if criteria:
            items.append({
                "criteria": criteria,
                "commentary": commentary
            })
    return items


def get_sheet_names(excel_path: str):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    return wb.sheetnames


def remove_header_only_rows(items):
    cleaned = []
    header_words = {
        "financials and e-billing",
        "operational",
        "shipper/ consignee information",
        "brokerage",
        "customs brokerage",
        "prealert and reporting",
        "automation"
    }

    for item in items:
        c = item["criteria"].strip().lower()
        if c in header_words:
            continue
        cleaned.append(item)
    return cleaned


# =========================
# EXCEL PARSING
# =========================
def parse_TEST_excel(excel_path: str, sheet_name: str):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name]

    business_start_date = safe_str(ws.cell(7, 3).value)

    mode_map = {
        "AFR": "Air Freight",
        "OFR": "Ocean Freight",
        "RFR": "Road Freight",
        "CDZ": "Customs"
    }

    selected_modes = []
    for r in range(8, 12):
        code = safe_str(ws.cell(r, 2).value)   # B
        tick = ws.cell(r, 3).value             # C
        if code and is_ticked(tick):
            selected_modes.append({
                "code": code,
                "name": mode_map.get(code, code)
            })

    selected_flows = []
    for r in range(12, 14):
        flow_label = safe_str(ws.cell(r, 2).value)  # B
        tick = ws.cell(r, 3).value                  # C
        if flow_label and is_ticked(tick):
            selected_flows.append(flow_label)

    finance_billing = remove_header_only_rows(read_qa_block(ws, 16, 22, q_col=2, a_col=4))
    operational = remove_header_only_rows(read_qa_block(ws, 26, 30, q_col=2, a_col=4))
    shipper_consignee = remove_header_only_rows(read_qa_block(ws, 33, 36, q_col=2, a_col=4))
    brokerage = remove_header_only_rows(read_qa_block(ws, 39, 42, q_col=2, a_col=4))
    customs_brokerage = remove_header_only_rows(read_qa_block(ws, 45, 48, q_col=2, a_col=4))
    prealert_reporting = remove_header_only_rows(read_qa_block(ws, 50, 53, q_col=2, a_col=4))
    automation = remove_header_only_rows(read_qa_block(ws, 56, 59, q_col=2, a_col=4))

    parsed = {
        "business_start_date": business_start_date,
        "modes": selected_modes,
        "flows": selected_flows,
        "finance_billing": finance_billing,
        "operational": operational,
        "shipper_consignee": shipper_consignee,
        "brokerage": brokerage,
        "customs_brokerage": customs_brokerage,
        "prealert_reporting": prealert_reporting,
        "automation": automation,
    }

    return parsed


# =========================
# PROMPT BUILDING
# =========================
def build_sop_prompt(parsed: dict) -> str:
    mode_text = ", ".join([m["name"] for m in parsed["modes"]]) if parsed["modes"] else "Not explicitly readable from Excel cell values"
    flow_text = ", ".join(parsed["flows"]) if parsed["flows"] else "Not explicitly readable from Excel cell values"

    return f"""
You are an expert logistics SOP writer for freight forwarding, customs, billing, and customer implementation handovers.

Your task is to write a complete, professional, detailed Standard Operating Procedure based on the extracted implementation checklist data below.

IMPORTANT:
- Use all criteria and all commentaries as real business requirements.
- Expand them into a usable operational SOP.
- If a commentary is blank, write a practical operational assumption and clearly label it as "Assumption:".
- If selected transport modes or import/export scope are not readable from the Excel cell values, mention that they must be confirmed during implementation kickoff.

EXTRACTED DATA:
{json.dumps(parsed, indent=2, ensure_ascii=False)}

READABLE SUMMARY:
- Business start date: {parsed["business_start_date"]}
- Selected transport modes: {mode_text}
- Selected import/export scope: {flow_text}

WRITE ONLY THIS SOP SECTION:
2. Operational Process (End-to-End)

MANDATORY STRUCTURE:
2.1 Objective
2.2 Scope
2.3 Transport Modes Covered
2.4 Import / Export Scope
2.5 Roles and Responsibilities
2.6 Detailed End-to-End Operational Process
2.7 Finance and Billing Requirements
2.8 Operational Requirements
2.9 Shipper / Consignee Information Requirements
2.10 Brokerage and Customs Brokerage Requirements
2.11 Prealert and Reporting Requirements
2.12 Automation Requirements
2.13 Exceptions and Escalation Process
2.14 Controls, Records, and Systems Used

WRITING RULES:
- Use clear business English.
- Use numbered operational steps where possible.
- Include If / Then logic where relevant.
- Use freight forwarding and customs wording.
- Make it detailed and realistic.
- Do not output JSON.
- Do not output explanations outside the SOP itself.
""".strip()


# =========================
# OLLAMA
# =========================
def check_ollama():
    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=10)
        r.raise_for_status()
        return True, "Ollama is reachable."
    except Exception as e:
        return False, f"Ollama is not reachable: {e}"


def warmup_model(model_name: str):
    payload = {
        "model": model_name,
        "prompt": "Reply only with OK.",
        "stream": False,
        "keep_alive": "30m"
    }
    r = requests.post(OLLAMA_URL, json=payload, timeout=(10, 180))
    r.raise_for_status()
    data = r.json()
    return data.get("response", "").strip()


def ollama_generate(prompt: str, model_name: str):
    payload = {
        "model": model_name,
        "prompt": prompt,
        "stream": False,
        "keep_alive": "30m",
        "options": {
            "temperature": 0.2
        }
    }

    r = requests.post(OLLAMA_URL, json=payload, timeout=(15, 600))
    r.raise_for_status()
    data = r.json()
    return data.get("response", "").strip()


# =========================
# WORD OUTPUT
# =========================
def insert_sop_into_template(template_path: str, sop_text: str, output_path: str):
    doc = Document(template_path)

    doc.add_page_break()

    p = doc.add_paragraph()
    r = p.add_run("2. Operational Process (End-to-End)")
    r.bold = True

    for line in sop_text.splitlines():
        clean_line = line.strip()

        if not clean_line:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        r = p.add_run(clean_line)

        if clean_line.startswith("2."):
            r.bold = True

    doc.save(output_path)


# =========================
# STREAMLIT UI
# =========================
st.set_page_config(page_title="SOP Generator", layout="wide")
st.title("SOP Generator with Excel + Word Template + Ollama")

st.write("Upload the Excel handover file and the Word template, then generate the SOP.")

model_name = st.text_input("Ollama model name", value=DEFAULT_MODEL)

excel_file = st.file_uploader(
    "Upload Excel file (.xlsx / .xlsm)",
    type=["xlsx", "xlsm"]
)

word_template = st.file_uploader(
    "Upload Word template (.docx)",
    type=["docx"]
)

sheet_name_input = st.text_input(
    "Excel sheet name",
    value=DEFAULT_SHEET_NAME
)

generate_btn = st.button("Generate SOP")


if generate_btn:
    if not excel_file:
        st.error("Please upload the Excel file.")
        st.stop()

    if not word_template:
        st.error("Please upload the Word template.")
        st.stop()

    with tempfile.TemporaryDirectory() as tmpdir:
        excel_path = os.path.join(tmpdir, excel_file.name)
        word_path = os.path.join(tmpdir, word_template.name)
        output_docx = os.path.join(tmpdir, "Generated_SOP.docx")

        with open(excel_path, "wb") as f:
            f.write(excel_file.getbuffer())

        with open(word_path, "wb") as f:
            f.write(word_template.getbuffer())

        ok, msg = check_ollama()
        if not ok:
            st.error(msg)
            st.stop()
        else:
            st.success(msg)

        try:
            sheet_names = get_sheet_names(excel_path)
            st.info(f"Sheets found in Excel: {', '.join(sheet_names)}")
        except Exception as e:
            st.error(f"Could not read Excel sheets: {e}")
            st.stop()

        if sheet_name_input not in sheet_names:
            st.error(
                f"Sheet '{sheet_name_input}' not found. Available sheets: {', '.join(sheet_names)}"
            )
            st.stop()

        try:
            with st.spinner("Parsing Excel file..."):
                parsed = parse_TEST_excel(excel_path, sheet_name_input)

            st.subheader("Parsed Excel Data")
            st.json(parsed)

            if not parsed["modes"]:
                st.warning("No transport modes were readable from column C. Your Excel likely uses visual checkboxes not linked to cells.")

            if not parsed["flows"]:
                st.warning("No Import/Export selection was readable from column C. Your Excel likely uses visual checkboxes not linked to cells.")

            with st.spinner("Building AI prompt..."):
                prompt = build_sop_prompt(parsed)

            with st.expander("Show AI Prompt"):
                st.text_area("Prompt", prompt, height=400)

            with st.spinner(f"Warming up Ollama model: {model_name} ..."):
                try:
                    warmup_result = warmup_model(model_name)
                    st.success(f"Warmup OK: {warmup_result}")
                except Exception as warmup_err:
                    st.warning(f"Warmup failed on {model_name}: {warmup_err}")
                    st.warning(f"Trying fallback model: {FALLBACK_MODEL}")
                    model_name = FALLBACK_MODEL
                    warmup_result = warmup_model(model_name)
                    st.success(f"Fallback warmup OK: {warmup_result}")

            with st.spinner(f"Generating SOP with local AI ({model_name})..."):
                try:
                    ai_text = ollama_generate(prompt, model_name)
                except Exception as gen_err:
                    st.warning(f"Generation failed with {model_name}: {gen_err}")
                    if model_name != FALLBACK_MODEL:
                        st.warning(f"Retrying with fallback model: {FALLBACK_MODEL}")
                        ai_text = ollama_generate(prompt, FALLBACK_MODEL)
                        model_name = FALLBACK_MODEL
                    else:
                        raise

            if not ai_text.strip():
                st.error("AI returned empty text.")
                st.stop()

            st.subheader("Generated SOP")
            st.text_area("SOP Output", ai_text, height=500)

            with st.spinner("Writing SOP into Word template..."):
                insert_sop_into_template(word_path, ai_text, output_docx)

            with open(output_docx, "rb") as f:
                st.download_button(
                    label="Download Generated SOP (.docx)",
                    data=f,
                    file_name="Generated_SOP.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            st.success("SOP generated successfully.")

        except Exception as e:
            st.error(f"Error: {e}")
